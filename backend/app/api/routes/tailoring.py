from io import BytesIO
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response
from sqlalchemy.orm import Session
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

from app.core.security import get_current_user
from app.db.session import get_db
from app.models.job import Job, JobMatch
from app.models.profile import CareerProfile
from app.models.resume import Resume
from app.models.tailoring import TailoredResume
from app.models.user import User
from app.schemas.tailoring import TailoringRequest, CoverLetterRequest
from app.services.resume_tailor import tailor_resume, generate_cover_letter


router = APIRouter(prefix="/api/tailoring", tags=["Tailoring"])


def _owned_profile(db: Session, user_id: int, profile_id: int):
    profile = (
        db.query(CareerProfile)
        .filter(
            CareerProfile.id == profile_id,
            CareerProfile.user_id == user_id,
        )
        .first()
    )
    if not profile:
        raise HTTPException(status_code=404, detail="Profile not found")
    return profile


def _serialize(item: TailoredResume):
    return {
        "id": item.id,
        "profile_id": item.profile_id,
        "resume_id": item.resume_id,
        "job_id": item.job_id,
        "version_name": item.version_name,
        "ats_score": item.ats_score,
        "professional_summary": item.professional_summary,
        "tailored_text": item.tailored_text,
        "selected_evidence": item.selected_evidence or [],
        "matched_keywords": item.matched_keywords or [],
        "missing_keywords": item.missing_keywords or [],
        "recommendations": item.recommendations or [],
        "cover_letter": item.cover_letter,
        "created_at": item.created_at.isoformat(),
        "updated_at": item.updated_at.isoformat(),
    }


@router.get("/job/{job_id}")
def job_workspace(
    job_id: int,
    profile_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    profile = _owned_profile(db, user.id, profile_id)
    job = db.query(Job).filter(Job.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    match = (
        db.query(JobMatch)
        .filter(
            JobMatch.profile_id == profile.id,
            JobMatch.job_id == job.id,
        )
        .first()
    )
    resumes = (
        db.query(Resume)
        .filter(Resume.profile_id == profile.id)
        .order_by(Resume.is_primary.desc(), Resume.created_at.desc())
        .all()
    )
    versions = (
        db.query(TailoredResume)
        .filter(
            TailoredResume.profile_id == profile.id,
            TailoredResume.job_id == job.id,
            TailoredResume.user_id == user.id,
        )
        .order_by(TailoredResume.created_at.desc())
        .all()
    )

    return {
        "profile": {
            "id": profile.id,
            "name": profile.name,
            "priority_keywords": profile.priority_keywords or [],
        },
        "job": {
            "id": job.id,
            "title": job.title,
            "company": job.company,
            "location": job.location,
            "description": job.description,
            "url": job.url,
            "salary": job.salary,
            "source": job.source,
            "remote": job.remote,
        },
        "match": None if not match else {
            "score": match.score,
            "title_score": match.title_score,
            "keyword_score": match.keyword_score,
            "location_score": match.location_score,
            "resume_score": match.resume_score,
            "matched_keywords": match.matched_keywords or [],
            "missing_keywords": match.missing_keywords or [],
            "concerns": match.concerns or [],
            "explanation": match.explanation,
        },
        "resumes": [
            {
                "id": resume.id,
                "name": resume.name,
                "is_primary": resume.is_primary,
                "analysis_score": resume.analysis_score,
            }
            for resume in resumes
        ],
        "versions": [_serialize(item) for item in versions],
    }


@router.post("/generate")
def generate(
    body: TailoringRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    profile = _owned_profile(db, user.id, body.profile_id)
    job = db.query(Job).filter(Job.id == body.job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    resume_query = db.query(Resume).filter(Resume.profile_id == profile.id)
    if body.resume_id:
        resume_query = resume_query.filter(Resume.id == body.resume_id)
    else:
        resume_query = resume_query.order_by(
            Resume.is_primary.desc(),
            Resume.created_at.desc(),
        )
    resume = resume_query.first()
    if not resume:
        raise HTTPException(
            status_code=400,
            detail="Upload a resume before tailoring.",
        )

    result = tailor_resume(
        resume_text=resume.extracted_text,
        job_title=job.title,
        company=job.company,
        job_description=job.description,
        profile_keywords=profile.priority_keywords or [],
    )

    version_name = (
        body.version_name
        or f"{job.company} - {job.title}"
    )[:255]

    item = TailoredResume(
        user_id=user.id,
        profile_id=profile.id,
        resume_id=resume.id,
        job_id=job.id,
        version_name=version_name,
        **result,
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    return _serialize(item)


@router.post("/{tailoring_id}/cover-letter")
def cover_letter(
    tailoring_id: int,
    body: CoverLetterRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    item = (
        db.query(TailoredResume)
        .filter(
            TailoredResume.id == tailoring_id,
            TailoredResume.user_id == user.id,
        )
        .first()
    )
    if not item:
        raise HTTPException(status_code=404, detail="Tailored resume not found")

    profile = db.query(CareerProfile).filter(CareerProfile.id == item.profile_id).first()
    job = db.query(Job).filter(Job.id == item.job_id).first()

    item.cover_letter = generate_cover_letter(
        candidate_name=profile.name if profile else "Candidate",
        job_title=job.title if job else "the position",
        company=job.company if job else "the company",
        location=job.location if job else "",
        evidence=item.selected_evidence or [],
        matched_keywords=item.matched_keywords or [],
    )
    db.commit()
    db.refresh(item)
    return _serialize(item)


@router.get("/versions/{profile_id}")
def versions(
    profile_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _owned_profile(db, user.id, profile_id)
    items = (
        db.query(TailoredResume)
        .filter(
            TailoredResume.profile_id == profile_id,
            TailoredResume.user_id == user.id,
        )
        .order_by(TailoredResume.created_at.desc())
        .all()
    )
    return [_serialize(item) for item in items]


@router.get("/{tailoring_id}")
def read_tailoring(
    tailoring_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    item = (
        db.query(TailoredResume)
        .filter(
            TailoredResume.id == tailoring_id,
            TailoredResume.user_id == user.id,
        )
        .first()
    )
    if not item:
        raise HTTPException(status_code=404, detail="Tailored resume not found")
    return _serialize(item)


def _owned_tailoring(db: Session, user_id: int, tailoring_id: int):
    item = (
        db.query(TailoredResume)
        .filter(
            TailoredResume.id == tailoring_id,
            TailoredResume.user_id == user_id,
        )
        .first()
    )
    if not item:
        raise HTTPException(status_code=404, detail="Tailored resume not found")
    return item


@router.get("/{tailoring_id}/download.txt")
def download_txt(
    tailoring_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    item = _owned_tailoring(db, user.id, tailoring_id)
    return Response(
        item.tailored_text,
        media_type="text/plain",
        headers={
            "Content-Disposition": f'attachment; filename="tailored_resume_{item.id}.txt"'
        },
    )


@router.get("/{tailoring_id}/download.docx")
def download_docx(
    tailoring_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    item = _owned_tailoring(db, user.id, tailoring_id)
    document = Document()
    document.add_heading(item.version_name, 0)
    document.add_heading("Professional Summary", level=1)
    document.add_paragraph(item.professional_summary)
    document.add_heading("Selected Relevant Evidence", level=1)
    for evidence in item.selected_evidence or []:
        document.add_paragraph(evidence, style="List Bullet")
    document.add_heading("Matched Keywords", level=1)
    document.add_paragraph(", ".join(item.matched_keywords or []))
    if item.cover_letter:
        document.add_page_break()
        document.add_heading("Cover Letter", level=1)
        for paragraph in item.cover_letter.split("\n\n"):
            document.add_paragraph(paragraph)

    stream = BytesIO()
    document.save(stream)
    return Response(
        stream.getvalue(),
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        headers={
            "Content-Disposition": f'attachment; filename="tailored_resume_{item.id}.docx"'
        },
    )


@router.get("/{tailoring_id}/download.pdf")
def download_pdf(
    tailoring_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    item = _owned_tailoring(db, user.id, tailoring_id)
    stream = BytesIO()
    styles = getSampleStyleSheet()
    story = [
        Paragraph(item.version_name, styles["Title"]),
        Spacer(1, 12),
        Paragraph("Professional Summary", styles["Heading1"]),
        Paragraph(item.professional_summary, styles["BodyText"]),
        Spacer(1, 12),
        Paragraph("Selected Relevant Evidence", styles["Heading1"]),
    ]
    for evidence in item.selected_evidence or []:
        story.append(Paragraph(f"- {evidence}", styles["BodyText"]))
        story.append(Spacer(1, 5))
    story.extend([
        Spacer(1, 12),
        Paragraph("Matched Keywords", styles["Heading1"]),
        Paragraph(", ".join(item.matched_keywords or []), styles["BodyText"]),
    ])
    if item.cover_letter:
        story.extend([
            Spacer(1, 18),
            Paragraph("Cover Letter", styles["Heading1"]),
        ])
        for paragraph in item.cover_letter.split("\n\n"):
            story.append(Paragraph(paragraph.replace("\n", "<br/>"), styles["BodyText"]))
            story.append(Spacer(1, 8))

    SimpleDocTemplate(stream, pagesize=letter).build(story)
    return Response(
        stream.getvalue(),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="tailored_resume_{item.id}.pdf"'
        },
    )
